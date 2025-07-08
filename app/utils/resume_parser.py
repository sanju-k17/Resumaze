"""
Resume Parser - Extract text from PDF and DOCX files
"""

import fitz  # PyMuPDF
from docx import Document
import os
import re

class ResumeParser:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def extract_text(self, file_path):
        """Extract text from resume file"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            
            doc.close()
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX using python-docx"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    def _clean_text(self, text):
        """Clean and normalize extracted text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\@\-$$$$\+]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    def extract_contact_info(self, text):
        """Extract contact information from resume text"""
        contact_info = {}
        
        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Phone extraction
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?$$?\d{3}$$?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact_info['phone'] = ''.join(phones[0]) if isinstance(phones[0], tuple) else phones[0]
        
        # LinkedIn extraction
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin = re.search(linkedin_pattern, text.lower())
        if linkedin:
            contact_info['linkedin'] = linkedin.group()
        
        return contact_info
    
    def extract_sections(self, text):
        """Extract different sections from resume"""
        sections = {}
        
        # Common section headers
        section_patterns = {
            'experience': r'(work experience|professional experience|employment|experience)',
            'education': r'(education|academic background|qualifications)',
            'skills': r'(skills|technical skills|core competencies|expertise)',
            'projects': r'(projects|personal projects|key projects)',
            'certifications': r'(certifications|certificates|licenses)'
        }
        
        text_lower = text.lower()
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text_lower)
            if match:
                start_pos = match.start()
                # Find next section or end of text
                next_section_pos = len(text)
                for other_pattern in section_patterns.values():
                    if other_pattern != pattern:
                        next_match = re.search(other_pattern, text_lower[start_pos + 50:])
                        if next_match:
                            next_section_pos = min(next_section_pos, start_pos + 50 + next_match.start())
                
                sections[section_name] = text[start_pos:next_section_pos].strip()
        
        return sections
